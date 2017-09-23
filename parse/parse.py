from docx import Document
import re
import json

document = Document('./data2.docx')

jsonData = {
  'questions': [
    # {
    #   'id': 1
    #   'text': 'test text',
    #   'variants': [{'id': 1, 'text': 'first'}, {'id': 2, 'text': 'second'}, {'id': 3, 'text': 'третий'}, {'id': 4, 'text': 'четвёртый'}],
    #   'answer': 1
    # }
  ]
}

reTitle = re.compile('^[\d]+.\s\[T\w+\d+\]\s')
reAnswer = re.compile('^[А|Б|В|Г]\)\s')

for idx, parargraph in enumerate(document.paragraphs):
  if re.match('^[\d]+.*', parargraph.text):
    
    print ("\rProcessing question: {}/{} ({}%)".format(idx+1, len(document.paragraphs), int(((idx+1)*100)/len(document.paragraphs))), end="")
    
    jsonData['questions'].append(
      {
        'id': idx+1,
        'text': reTitle.sub('', parargraph.text).lower().capitalize(),
        'variants': [
          { 'id': 1, 'text': reAnswer.sub('', document.paragraphs[idx + 1].text).capitalize() },
          { 'id': 2, 'text': reAnswer.sub('', document.paragraphs[idx + 2].text).capitalize() },
          { 'id': 3, 'text': reAnswer.sub('', document.paragraphs[idx + 3].text).capitalize() },
          { 'id': 4, 'text': reAnswer.sub('', document.paragraphs[idx + 4].text).capitalize() }
        ],
        'answer': 1
      }
    )

outputFile = open("jsonData.json", "w")
outputFile.write(json.dumps(jsonData, sort_keys=True, indent=2, ensure_ascii=False))
outputFile.close()
